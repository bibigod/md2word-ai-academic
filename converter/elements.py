"""
元素创建函数：段落、标题、表格、检测函数等
"""
import re
import math

from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .config import DEFAULT_CONFIG


# ========== 工具函数 ==========

def set_run_font(run, cn_font='仿宋_GB2312', size=12, bold=False,
                  en_font='Times New Roman'):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)


def make_para(doc, text, cn_font='仿宋_GB2312', size=12, bold=False,
              en_font='Times New Roman',
              align=None, indent=True, space_before=0, space_after=0):
    p = doc.add_paragraph()
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, cn_font, size, bold, en_font=en_font)
    return p


def make_heading_para(doc, text, level, cn_font='黑体', size=16, bold=True,
                      en_font='Times New Roman'):
    """
    创建使用 Word 内置 Heading 样式的标题段落。
    level: 1-6，对应 Heading 1 ~ Heading 6
    """
    p = doc.add_heading('', level=level)
    for r in p.runs:
        r.text = ''
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, cn_font, size, bold, en_font=en_font)
    return p


def _calc_text_width(text):
    """估算文本显示宽度：中文/全角算2，其余算1"""
    w = 0
    for ch in text:
        if '\u4e00' <= ch <= '\u9fff' or '\uff00' <= ch <= '\uffef' or '\u3000' <= ch <= '\u303f':
            w += 2
        else:
            w += 1
    return w


def _set_col_widths(table, widths_cm):
    """通过 tblGrid/gridCol + tcW 设置固定列宽"""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()

    # 固定布局
    for old in tbl_pr.findall(qn('w:tblLayout')):
        tbl_pr.remove(old)
    layout = tbl_pr.makeelement(qn('w:tblLayout'), {qn('w:type'): 'fixed'})
    tbl_pr.append(layout)

    # 重写 tblGrid
    old_grid = tbl.find(qn('w:tblGrid'))
    if old_grid is not None:
        tbl.remove(old_grid)
    new_grid = tbl.makeelement(qn('w:tblGrid'), {})
    for w_cm in widths_cm:
        twips = int(w_cm * 360000 / 635)
        col_el = new_grid.makeelement(qn('w:gridCol'), {qn('w:w'): str(twips)})
        new_grid.append(col_el)
    tbl_pr.addnext(new_grid)

    # 每个单元格 tcW
    for row in table.rows:
        for j, w_cm in enumerate(widths_cm):
            tc = row.cells[j]._tc
            tc_pr = tc.tcPr if tc.tcPr is not None else tc._add_tcPr()
            for old_w in tc_pr.findall(qn('w:tcW')):
                tc_pr.remove(old_w)
            twips = int(w_cm * 360000 / 635)
            tcW = tc_pr.makeelement(qn('w:tcW'), {
                qn('w:w'): str(twips), qn('w:type'): 'dxa',
            })
            tc_pr.insert(0, tcW)


def _text_width_cm(text, font_size_pt):
    """
    根据字号计算文本渲染宽度（cm）。
    原理：中文字体是方块字，字宽 = 字号（这是字体设计标准，换任何中文字体都成立）。
    英文/数字平均宽度约为字号的 0.5 倍。
    """
    PT_TO_CM = 0.03528  # 1pt = 0.03528cm (国际标准)
    w = 0.0
    for ch in text:
        if '\u4e00' <= ch <= '\u9fff' or '\uff00' <= ch <= '\uffef' or '\u3000' <= ch <= '\u303f':
            w += font_size_pt * PT_TO_CM       # 全角：宽 = 字号
        else:
            w += font_size_pt * PT_TO_CM * 0.5  # 半角：约一半
    return w


# Word 默认单元格内边距：左 0.19cm + 右 0.19cm
CELL_PAD_CM = 0.38


def _calc_col_widths(headers, rows, header_size_pt=11, body_size_pt=11,
                     avail_cm=14.66):
    """
    计算各列宽度（cm）。

    硬约束：任何单元格最多 2 行，绝不出现 3 行。
    优先级：能单行的尽量单行 → 不行就 2 行 → 压缩其他列来保证不超 2 行。

    min_cm   = 表头单行宽度（表头绝不换行）
    two_cm   = 内容最多 2 行所需的最小宽度（硬底线）
    ideal_cm = 内容单行宽度（最优目标）
    """
    n_cols = len(headers)
    PT_TO_CM = 0.03528

    # 渲染安全余量：半个中文字符宽度（随字号缩放）
    safety = body_size_pt * PT_TO_CM * 0.5

    # --- 计算三档宽度 ---

    # 表头单行宽度
    header_cm = [_text_width_cm(h, header_size_pt) + CELL_PAD_CM + safety
                 for h in headers]

    # 内容单行宽度（理想）
    content_cm = [0.0] * n_cols
    for row_data in rows:
        for j in range(min(len(row_data), n_cols)):
            need = _text_width_cm(row_data[j], body_size_pt) + CELL_PAD_CM + safety
            content_cm[j] = max(content_cm[j], need)

    ideal_cm = [max(header_cm[j], content_cm[j]) for j in range(n_cols)]

    # 内容最多 2 行的最小宽度：文字部分折半 + 内边距 + 余量
    # 即每行至少放 ceil(总字数/2) 个字的宽度
    # 必须向上取整到整字宽度，否则 Word 按整字断行会多出一行
    char_w = body_size_pt * PT_TO_CM  # 一个中文字符的宽度（最大字符单位）
    two_line_cm = [0.0] * n_cols
    for j in range(n_cols):
        content_text = content_cm[j] - CELL_PAD_CM - safety  # 纯文字宽度
        # 向上取整到整字宽度：保证每行能放下 ceil(总字数/2) 个字
        chars_per_line = math.ceil(content_text / 2.0 / char_w) if char_w > 0 else 1
        half_text = chars_per_line * char_w
        two_line_cm[j] = half_text + CELL_PAD_CM + safety

    # 最终 min = max(表头单行, 内容两行)
    min_cm = [max(header_cm[j], two_line_cm[j]) for j in range(n_cols)]

    total_ideal = sum(ideal_cm)
    total_min = sum(min_cm)

    # 情况1：所有列都能单行
    if total_ideal <= avail_cm:
        leftover = avail_cm - total_ideal
        bonus = leftover / n_cols
        return [ideal_cm[j] + bonus for j in range(n_cols)]

    # 情况2：连 2 行底线都放不下，等比缩放（极端情况）
    if total_min >= avail_cm:
        ratio = avail_cm / total_min
        return [w * ratio for w in min_cm]

    # 情况3：部分列单行，部分列 2 行
    # 贪心：从最窄列开始，尝试锁定为单行
    indices = list(range(n_cols))
    indices.sort(key=lambda j: ideal_cm[j])

    result = [0.0] * n_cols
    locked = [False] * n_cols
    remaining_space = avail_cm

    for j in indices:
        # 锁定此列为单行后，剩余空间是否够给其他列的 2 行底线
        other_min = sum(min_cm[k] for k in range(n_cols) if not locked[k] and k != j)
        if remaining_space - ideal_cm[j] >= other_min:
            result[j] = ideal_cm[j]
            locked[j] = True
            remaining_space -= ideal_cm[j]
        else:
            break

    # 未锁定列：按理想宽度比例分配剩余空间（但不低于 2 行底线）
    unlocked = [j for j in range(n_cols) if not locked[j]]
    if unlocked:
        unlocked_ideal_total = sum(ideal_cm[j] for j in unlocked)
        for j in unlocked:
            if unlocked_ideal_total > 0:
                share = ideal_cm[j] / unlocked_ideal_total * remaining_space
            else:
                share = remaining_space / len(unlocked)
            result[j] = max(share, min_cm[j])

    # 微调：从宽裕列借空间给"差一点就单行"的列
    unlocked_by_deficit = sorted(unlocked, key=lambda j: ideal_cm[j] - result[j])
    for j in unlocked_by_deficit:
        deficit = ideal_cm[j] - result[j]
        if deficit <= 0:
            continue
        # 从其他未锁定列借（从最宽的开始，最多借到它的 2 行底线）
        donors = sorted([k for k in unlocked if k != j and result[k] > min_cm[k]],
                        key=lambda k: result[k], reverse=True)
        for donor in donors:
            if deficit <= 0:
                break
            can_give = result[donor] - min_cm[donor]
            give = min(can_give, deficit)
            if give > 0:
                result[donor] -= give
                result[j] += give
                deficit -= give

    # 最终修正：确保总和不超出
    total_result = sum(result)
    if total_result > avail_cm * 1.001:
        ratio = avail_cm / total_result
        result = [r * ratio for r in result]

    return result


def _set_row_height(row, table_cfg):
    """根据配置设置表格行高"""
    if 'row_height_pt' in table_cfg:
        # 固定行高（磅）
        tr = row._tr
        tr_pr = tr.get_or_add_trPr()
        for old in tr_pr.findall(qn('w:trHeight')):
            tr_pr.remove(old)
        twips = int(table_cfg['row_height_pt'] * 20)  # 1pt = 20 twips
        tr_height = tr_pr.makeelement(qn('w:trHeight'), {
            qn('w:val'): str(twips),
            qn('w:hRule'): 'atLeast',  # 最小行高，内容多时可撑开
        })
        tr_pr.append(tr_height)
    elif 'row_height_line' in table_cfg:
        # 倍数行距 —— 设置到行内段落的行距上
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = table_cfg['row_height_line']


def _set_cell_border(cell, **kwargs):
    """
    设置单元格边框。
    用法: _set_cell_border(cell, top={"sz": 12, "val": "single"}, ...)
    sz 单位是 1/8 磅: sz=12 → 1.5pt, sz=6 → 0.75pt
    val="single" 实线, val="nil" 无线
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge, data in kwargs.items():
        tag = f'w:{edge}'
        el = tcBorders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tcBorders.append(el)
        for attr in ('sz', 'val', 'color', 'space'):
            if attr in data:
                el.set(qn(f'w:{attr}'), str(data[attr]))


# 三线表边框预设
_BORDER_NIL = {"val": "nil"}
_BORDER_THICK = {"sz": "12", "val": "single", "color": "000000", "space": "0"}  # 1.5pt
_BORDER_THIN = {"sz": "6", "val": "single", "color": "000000", "space": "0"}    # 0.75pt


def _apply_three_line_style(table, n_rows):
    """
    三线表：顶线1.5pt + 栏目线0.75pt + 底线1.5pt，无竖线。
    符合 GB/T 7713.1-2006 学术论文表格规范。
    """
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            borders = {
                "start": _BORDER_NIL,     # 无左竖线
                "end": _BORDER_NIL,       # 无右竖线
            }
            # 顶线（仅表头行顶部）
            if i == 0:
                borders["top"] = _BORDER_THICK       # 顶线 1.5pt
                borders["bottom"] = _BORDER_THIN     # 栏目线 0.75pt
            elif i == n_rows - 1:
                borders["top"] = _BORDER_NIL
                borders["bottom"] = _BORDER_THICK    # 底线 1.5pt
            else:
                borders["top"] = _BORDER_NIL
                borders["bottom"] = _BORDER_NIL
            _set_cell_border(cell, **borders)


def add_table_with_data(doc, headers, rows, cfg=None):
    if cfg is None:
        cfg = DEFAULT_CONFIG
    table_cfg = cfg.get('table', {})

    h_font = table_cfg.get('header_font_cn', '黑体')
    h_size = table_cfg.get('header_font_size_pt', 11)
    h_bold = table_cfg.get('header_bold', True)
    b_font = table_cfg.get('body_font_cn', '仿宋_GB2312')
    b_size = table_cfg.get('body_font_size_pt', 11)
    en_font = cfg.get('body', {}).get('font_en', 'Times New Roman')
    table_style = table_cfg.get('style', 'grid')

    n_cols = len(headers)
    avail_cm = 21.0 - cfg['page']['left_margin_cm'] - cfg['page']['right_margin_cm']
    col_widths_cm = _calc_col_widths(headers, rows,
                                     header_size_pt=h_size,
                                     body_size_pt=b_size,
                                     avail_cm=avail_cm)

    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _set_col_widths(table, col_widths_cm)

    # 应用表格边框样式
    if table_style == 'three_line':
        _apply_three_line_style(table, 1 + len(rows))

    # 表头行
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(h), h_font, h_size, h_bold, en_font=en_font)
    _set_row_height(table.rows[0], table_cfg)

    # 数据行
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(val), b_font, b_size, en_font=en_font)
        _set_row_height(table.rows[i + 1], table_cfg)

    # 表后空行
    make_para(doc, '', indent=False)


def parse_table_row(line):
    """解析 Markdown 表格行，返回单元格列表"""
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [cell.strip() for cell in line.split('|')]


def is_separator_row(line):
    """判断是否为表格分隔行，如 |---|---|---|"""
    return bool(re.match(r'^\|?\s*[-:]+[-|\s:]*$', line.strip()))


def is_table_row(line):
    """判断是否为表格行：要求行首尾都是 | 或至少有 2 个 | 分隔"""
    stripped = line.strip()
    if stripped.startswith('#'):
        return False
    # 标准 Markdown 表格行以 | 开头或结尾
    if stripped.startswith('|') or stripped.endswith('|'):
        return stripped.count('|') >= 2
    return False


_CN_NUMS = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
            '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十']


def _to_chinese_num(n):
    """阿拉伯数字转中文（1~20）"""
    if 1 <= n <= len(_CN_NUMS):
        return _CN_NUMS[n - 1]
    return str(n)


def _strip_manual_numbering(text):
    """去掉标题文本中的手动编号，如 '（1）信号特征提取' → '信号特征提取'
    支持的编号格式:
      - (1) （1） (1.2) 等括号编号
      - 4.2.2 、1.1 等数字点号编号（后跟空格或中文）
      - 1 、2 等纯数字编号（后跟空格或中文）
      - 第一章、第1章 等章节编号
    """
    # 匹配 (1) （1） (1.2) 等开头的编号
    m = re.match(r'^[（(]\s*[\d.]+\s*[）)]\s*', text)
    if m:
        return text[m.end():]
    # 匹配 "第x章" / "第一章" 等（x可以是数字或中文数字）
    m = re.match(r'^第[一二三四五六七八九十\d]+章\s*', text)
    if m:
        return text[m.end():]
    # 匹配 "4.2.2 " / "1.1 " 等数字点号编号（至少一个点，后跟空格或中文）
    m = re.match(r'^(\d+\.)+\d+(?=\s|[\u4e00-\u9fff])\s*', text)
    if m:
        return text[m.end():]
    # 匹配纯数字编号 "1 " / "2 "（后跟空格或中文，但不匹配整个文本就是数字的情况）
    m = re.match(r'^\d+(?=\s+[\S]|[\u4e00-\u9fff])\s*', text)
    if m and m.end() < len(text):
        return text[m.end():]
    return text


def detect_heading(line):
    """检测标题级别，返回 (level, text) 或 None"""
    m = re.match(r'^(#{1,6})\s+(.+)$', line)
    if m:
        return len(m.group(1)), m.group(2)
    return None


def detect_fig_caption(line):
    """检测图标题行：以'图'开头，返回标题文字部分（去掉手动编号），或 None"""
    stripped = line.strip()
    m = re.match(r'^图\s*[\d.]*\s*(.+)$', stripped)
    if m:
        return m.group(1).strip()
    return None


def detect_table_caption(line):
    """检测表标题行：以'表'开头，返回标题文字部分（去掉手动编号），或 None"""
    stripped = line.strip()
    m = re.match(r'^表\s*[\d.]*\s*(.+)$', stripped)
    if m:
        return m.group(1).strip()
    return None


def detect_formula(line):
    """
    检测公式行：独占一行且包含 $...$ 的行，可带序号 (n) 或 (x.y)。
    返回 (formula_text, has_label) 或 None。
    formula_text 包含 $ 符号（保留给 MathType 转换）。
    """
    stripped = line.strip()
    # 匹配: $...$ 后面可选跟 (序号)，序号可以是 (1) 或 (1.2) 等
    m = re.match(r'^(\$[^$]+\$)\s*(\([0-9.]+\))?\s*$', stripped)
    if m:
        return m.group(1), m.group(2) is not None
    return None


def make_formula_para(doc, formula_text, label, cfg):
    """
    创建公式段落：使用"公式"样式（制表位已在样式中定义）。
    布局: [Tab]公式[Tab](n)
    """
    body_cfg = cfg['body']
    en_font = body_cfg.get('font_en', 'Times New Roman')

    p = doc.add_paragraph(style='公式')

    run_tab1 = p.add_run('\t')
    set_run_font(run_tab1, body_cfg['font_cn'], body_cfg['font_size_pt'], en_font=en_font)

    run_formula = p.add_run(formula_text)
    set_run_font(run_formula, body_cfg['font_cn'], body_cfg['font_size_pt'], en_font=en_font)

    if label:
        run_tab2 = p.add_run('\t')
        set_run_font(run_tab2, body_cfg['font_cn'], body_cfg['font_size_pt'], en_font=en_font)
        run_label = p.add_run(label)
        set_run_font(run_label, body_cfg['font_cn'], body_cfg['font_size_pt'], en_font=en_font)

    return p


def detect_unordered_list(line):
    """检测无序列表行: - item / * item / + item，返回 (indent_level, text) 或 None"""
    m = re.match(r'^(\s*)[*\-+]\s+(.+)$', line)
    if m:
        indent = len(m.group(1))
        level = indent // 2  # 每2个空格算一级缩进
        return level, m.group(2)
    return None


def detect_ordered_list(line):
    """检测有序列表行: 1. item / 2. item，返回 (indent_level, text) 或 None"""
    m = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
    if m:
        indent = len(m.group(1))
        level = indent // 2
        return level, m.group(2)
    return None
