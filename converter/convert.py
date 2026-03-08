"""
主转换逻辑
"""
import sys
import os

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .config import DEFAULT_CONFIG, load_config
from .elements import (
    set_run_font, make_para, make_heading_para,
    add_table_with_data, parse_table_row, is_separator_row, is_table_row,
    detect_heading, detect_fig_caption, detect_table_caption, detect_formula,
    make_formula_para, detect_unordered_list, detect_ordered_list,
    _to_chinese_num, _strip_manual_numbering,
)
from .numbering import _new_num_instance, make_list_para
from .styles import setup_document_styles


# ========== 主转换逻辑 ==========

def convert(md_path, docx_path=None, cfg=None):
    if cfg is None:
        cfg = DEFAULT_CONFIG
    if docx_path is None:
        docx_path = os.path.splitext(md_path)[0] + '.docx'

    body_cfg = cfg['body']
    page_cfg = cfg['page']

    doc = Document()

    # 设置所有样式并获取编号 ID
    list_num_id = setup_document_styles(doc, cfg)

    for section in doc.sections:
        section.top_margin = Cm(page_cfg['top_margin_cm'])
        section.bottom_margin = Cm(page_cfg['bottom_margin_cm'])
        section.left_margin = Cm(page_cfg['left_margin_cm'])
        section.right_margin = Cm(page_cfg['right_margin_cm'])

    # ===== 页眉页脚 =====
    hf_cfg = cfg.get('header_footer', {})
    header_text = hf_cfg.get('header_text', '')
    page_number = hf_cfg.get('page_number', False)

    for section in doc.sections:
        # --- 页眉 ---
        header = section.header
        header.is_linked_to_previous = False
        if header_text:
            p = header.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header_text)
            set_run_font(run,
                         hf_cfg.get('header_font_cn', '仿宋_GB2312'),
                         hf_cfg.get('header_font_size_pt', 9))
            # 页眉底部边框线
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')       # 0.75pt
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '000000')
            pBdr.append(bottom)
            pPr.append(pBdr)

        # --- 页脚页码 ---
        footer = section.footer
        footer.is_linked_to_previous = False
        if page_number:
            p = footer.paragraphs[0]
            pn_position = hf_cfg.get('page_number_position', 'bottom_center')
            if pn_position == 'bottom_right':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            pn_font_cn = hf_cfg.get('header_font_cn', '仿宋_GB2312')
            pn_font_size = hf_cfg.get('page_number_font_size_pt', 9)

            # PAGE 域代码: begin
            run1 = p.add_run()
            set_run_font(run1, pn_font_cn, pn_font_size)
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run1._element.append(fldChar1)

            # PAGE 域代码: instrText
            run2 = p.add_run()
            set_run_font(run2, pn_font_cn, pn_font_size)
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = ' PAGE '
            run2._element.append(instrText)

            # PAGE 域代码: end
            run3 = p.add_run()
            set_run_font(run3, pn_font_cn, pn_font_size)
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run3._element.append(fldChar2)

    # 读取文件
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # 编号状态
    formula_cfg = cfg.get('formula', {})
    formula_numbering = formula_cfg.get('numbering', 'plain')  # plain | chapter
    chapter_num = 0       # 当前章节号（h1 递增）
    formula_counter = 0   # 当前章内公式序号
    fig_counter = 0       # 当前章内图序号
    table_counter = 0     # 当前章内表序号
    cap_cfg = cfg.get('caption', {})

    # 标题自动编号状态
    auto_numbering = cfg['headings'].get('auto_numbering', False)
    h1_counter = 0
    h2_counter = 0
    h3_counter = 0
    h4_counter = 0

    # 列表状态
    in_list = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # 标题
        heading = detect_heading(line)
        if heading:
            level, text = heading
            # h1 标题触发章节号递增，所有序号重置
            if level == 1:
                if formula_numbering == 'chapter':
                    chapter_num += 1
                    formula_counter = 0
                    fig_counter = 0
                    table_counter = 0

            # 自动编号：更新计数器并生成编号前缀
            if auto_numbering:
                text = _strip_manual_numbering(text)
                if level == 1:
                    h1_counter += 1
                    h2_counter = 0
                    h3_counter = 0
                    h4_counter = 0
                    prefix = f'{_to_chinese_num(h1_counter)}、'
                elif level == 2:
                    h2_counter += 1
                    h3_counter = 0
                    h4_counter = 0
                    prefix = f'{h1_counter}.{h2_counter} '
                elif level == 3:
                    h3_counter += 1
                    h4_counter = 0
                    prefix = f'{h1_counter}.{h2_counter}.{h3_counter} '
                elif level == 4:
                    h4_counter += 1
                    prefix = f'{h1_counter}.{h2_counter}.{h3_counter}.{h4_counter} '
                else:
                    prefix = ''
                text = prefix + text

            h_key = f'h{level}'
            h_cfg = cfg['headings'].get(h_key, {'font_cn': '黑体', 'size_pt': 12, 'bold': True})
            # 使用 Word 内置 Heading 样式，用户可在 Word 中统一修改
            make_heading_para(doc, text, level, h_cfg['font_cn'],
                              h_cfg['size_pt'], h_cfg.get('bold', True),
                              en_font=body_cfg.get('font_en', 'Times New Roman'))
            i += 1
            continue

        # 公式行：$...$ 可带序号
        formula = detect_formula(line)
        if formula:
            formula_text, has_label = formula
            # 自动生成编号
            label = None
            if has_label:
                formula_counter += 1
                if formula_numbering == 'chapter' and chapter_num > 0:
                    label = f'({chapter_num}.{formula_counter})'
                else:
                    label = f'({formula_counter})'
            make_formula_para(doc, formula_text, label, cfg)
            i += 1
            continue

        # 表标题行（"表xxx"）：记下标题，等下一块表格出现时插入到表格上方
        table_cap_text = detect_table_caption(line)
        if table_cap_text:
            # 先看后面是不是紧跟表格（跳过空行）
            peek = i + 1
            while peek < len(lines) and not lines[peek].strip():
                peek += 1
            if peek < len(lines) and is_table_row(lines[peek].rstrip()):
                # 生成编号
                table_counter += 1
                if formula_numbering == 'chapter' and chapter_num > 0:
                    cap_label = f'表{chapter_num}.{table_counter}'
                else:
                    cap_label = f'表{table_counter}'
                # 插入表标题（使用"表标题"样式）
                p = doc.add_paragraph(style='表标题')
                run = p.add_run(f'{cap_label}  {table_cap_text}')
                set_run_font(run, cap_cfg.get('font_cn', '黑体'),
                             cap_cfg.get('font_size_pt', 11), bold=True,
                             en_font=body_cfg.get('font_en', 'Times New Roman'))
                i += 1
                continue
            else:
                # 后面不是表格，当正文处理
                pass

        # 表格：收集连续的表格行
        if is_table_row(line):
            table_lines = []
            while i < len(lines) and is_table_row(lines[i].rstrip()):
                table_lines.append(lines[i].rstrip())
                i += 1

            # 过滤掉分隔行，解析表头和数据行
            headers = None
            rows = []
            for tl in table_lines:
                if is_separator_row(tl):
                    continue
                cells = parse_table_row(tl)
                if headers is None:
                    headers = cells
                else:
                    while len(cells) < len(headers):
                        cells.append('')
                    rows.append(cells[:len(headers)])

            if headers:
                add_table_with_data(doc, headers, rows, cfg)
            continue

        # 图标题行（"图xxx"）：图标题在图下方，黑体居中，自动编号
        fig_cap_text = detect_fig_caption(line)
        if fig_cap_text:
            fig_counter += 1
            if formula_numbering == 'chapter' and chapter_num > 0:
                cap_label = f'图{chapter_num}.{fig_counter}'
            else:
                cap_label = f'图{fig_counter}'
            # 使用"图标题"样式
            p = doc.add_paragraph(style='图标题')
            run = p.add_run(f'{cap_label}  {fig_cap_text}')
            set_run_font(run, cap_cfg.get('font_cn', '黑体'),
                         cap_cfg.get('font_size_pt', 11), bold=True,
                         en_font=body_cfg.get('font_en', 'Times New Roman'))
            i += 1
            continue

        # 无序列表 / 有序列表 → Word 内置编号
        ul = detect_unordered_list(line)
        ol = detect_ordered_list(line) if not ul else None
        if ul or ol:
            level, text = ul if ul else ol
            level = min(level, 3)  # 最多4级

            if not in_list:
                # 新列表块开始，创建新的编号实例（重新从1开始）
                list_num_id = _new_num_instance(doc, list_num_id)
                in_list = True

            make_list_para(doc, text, level, cfg, num_id=list_num_id)
            i += 1
            continue
        else:
            # 非列表行，重置列表状态
            if in_list:
                in_list = False

        # 普通正文
        make_para(doc, line.strip(),
                  cn_font=body_cfg['font_cn'], size=body_cfg['font_size_pt'],
                  en_font=body_cfg.get('font_en', 'Times New Roman'),
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        i += 1

    # 保存
    doc.save(docx_path)

    # 统计
    total_chars = sum(len(p.text) for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                total_chars += len(cell.text)
    print(f'输入文件: {md_path}')
    print(f'输出文件: {docx_path}')
    print(f'总字符数: {total_chars}')
    print(f'段落数:   {len(doc.paragraphs)}')
    print(f'表格数:   {len(doc.tables)}')
    print(f'估算页数: {total_chars / 1000:.1f} 页')
    print('转换完成!')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python -m converter.convert input.md [output.docx] [config.yaml]')
        sys.exit(1)
    md_file = sys.argv[1]
    out_file = sys.argv[2] if len(sys.argv) > 2 else None
    cfg_file = sys.argv[3] if len(sys.argv) > 3 else None
    cfg = load_config(cfg_file)
    convert(md_file, out_file, cfg)
